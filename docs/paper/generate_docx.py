"""Build an ApJS-style .docx for the Circex paper, modeled on Sharma et al. 2026."""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

REPO = Path("/Users/ericphillips/Documents/IntelliJ/Projects/IDEA/Circex")
FIG = REPO / "docs/images/eval_4way.png"
OUT = REPO / "docs/paper/Circex_paper.docx"

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(9.5)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0

for m in ("top", "bottom", "left", "right"):
    setattr(doc.sections[0], f"{m}_margin", Inches(0.75))


def _cols(section, n, space=0.3):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(n))
    cols.set(qn("w:space"), str(int(space * 1440)))


def para(text="", *, align=None, size=9.5, bold=False, italic=False, before=0, after=0,
         first_indent=0.2, justify=True, style=None):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_before, pf.space_after = Pt(before), Pt(after)
    if first_indent:
        pf.first_line_indent = Inches(first_indent)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.size, r.bold, r.italic = Pt(size), bold, italic
    return p


def heading(num, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{num}. {text}" if num else text)
    r.bold, r.font.size = True, Pt(size)
    return p


def subheading(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{num}. {text}")
    r.italic, r.font.size = True, Pt(9.5)
    return p


# ---------- footnotes (real Word footnotes; fall back to inline) ----------
_FN = {"id": 0, "part": None}


def _ensure_footnotes_part():
    if _FN["part"] is not None:
        return _FN["part"]
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    ct = ("application/vnd.openxmlformats-officedocument."
          "wordprocessingml.footnotes+xml")
    xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
        '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r>'
        '<w:continuationSeparator/></w:r></w:p></w:footnote>'
        '</w:footnotes>'
    ).encode()
    part = Part(PackURI("/word/footnotes.xml"), ct, xml, doc.part.package)
    doc.part.relate_to(
        part, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes")
    _FN["part"], _FN["id"] = part, 1
    return part


def footnote(paragraph, text):
    try:
        part = _ensure_footnotes_part()
        fid = _FN["id"]
        _FN["id"] += 1
        from lxml import etree
        root = etree.fromstring(part.blob)
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        fn = etree.SubElement(root, f"{{{W}}}footnote")
        fn.set(f"{{{W}}}id", str(fid))
        p = etree.SubElement(fn, f"{{{W}}}p")
        r0 = etree.SubElement(p, f"{{{W}}}r")
        etree.SubElement(etree.SubElement(r0, f"{{{W}}}rPr"), f"{{{W}}}rStyle").set(
            f"{{{W}}}val", "FootnoteReference")
        etree.SubElement(r0, f"{{{W}}}footnoteRef")
        r1 = etree.SubElement(p, f"{{{W}}}r")
        t = etree.SubElement(r1, f"{{{W}}}t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = " " + text
        part._blob = etree.tostring(root)
        # reference mark in the body
        run = paragraph.add_run()
        rpr = OxmlElement("w:rPr")
        va = OxmlElement("w:vertAlign"); va.set(qn("w:val"), "superscript"); rpr.append(va)
        run._r.append(rpr)
        ref = OxmlElement("w:footnoteReference"); ref.set(qn("w:id"), str(fid))
        run._r.append(ref)
    except Exception:
        paragraph.add_run(f" [{text}]").font.size = Pt(8)


def add_super(paragraph, text):
    r = paragraph.add_run(text)
    r.font.superscript = True
    r.font.size = Pt(7)


# ==================================================================
# TITLE BLOCK  (full width, one column)
# ==================================================================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(6)
rt = t.add_run("Schema-Constrained Structured Extraction of Optical Observations\n"
               "from GCN Circulars for Real-Time Ingestion into SkyPortal")
rt.bold, rt.font.size = True, Pt(15)

au = doc.add_paragraph(); au.alignment = WD_ALIGN_PARAGRAPH.CENTER
au.paragraph_format.space_after = Pt(4)
au.add_run("Eric Phillips").font.size = Pt(11)
add_super(au, "1")
au.add_run(", Sushant Sharma Chaudhary").font.size = Pt(11)
add_super(au, "1")
au.add_run(", and [additional authors]").font.size = Pt(11)
add_super(au, "1")

aff = doc.add_paragraph(); aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff.paragraph_format.space_after = Pt(2)
ra = aff.add_run("¹ School of Physics and Astronomy, University of Minnesota, "
                 "Minneapolis, MN 55455, USA; phill923@umn.edu")
ra.font.size, ra.italic = Pt(8), True

dt = doc.add_paragraph(); dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
dt.paragraph_format.space_after = Pt(8)
rd = dt.add_run("Draft compiled 2026 July 15")
rd.font.size, rd.italic = Pt(8), True

# ---- Abstract ----
ah = doc.add_paragraph(); ah.alignment = WD_ALIGN_PARAGRAPH.CENTER
ah.add_run("Abstract").bold = True
abstract = (
    "The General Coordinates Network (GCN) Circulars are human-written reports of "
    "time-domain and multimessenger observations; the archive of more than 40,500 "
    "prose Circulars is difficult to ingest into structured downstream systems. "
    "Building directly on the LLM-based analysis of the Circulars archive by "
    "Sharma et al. (2026, hereafter S25) — whose released topic labels, archive, "
    "and Swift-validated redshift table we adopt — we present Circex, a pipeline "
    "that converts optical Circulars into validated JSON conforming to a single "
    "output schema (nasa-gcn/gcn-schema). Extractors sharing that schema are "
    "compared on 500 rows of the S25 gold set: a transparent regular-expression "
    "baseline, and the same Mistral-7B-Instruct-v0.2 model used by S25, served "
    "locally with grammar-constrained decoding so the sampler cannot emit a token "
    "that violates the schema. The constrained model attains F1 = 0.935 on redshift "
    "extraction, against 0.690 for the identical model under free-form generation "
    "(Δ = +0.245) and 0.862 for the regex baseline. The improvement is confined "
    "to recall (0.559 → 0.937) at comparable precision (0.903 → 0.932): under "
    "free-form generation the model fails to emit parseable output that is then "
    "discarded, rather than extracting incorrectly. We conclude that the apparent "
    "weakness of open 7B models on this task is substantially an artifact of the "
    "output serialization rather than of the model, and that constrained decoding "
    "recovers extraction ability with no fine-tuning, retrieval, or larger model. "
    "The converse is also reported: on lexically regular event names the regex "
    "baseline (0.869) exceeds both language-model configurations, so the strongest "
    "system is a hybrid that routes each field to the extractor whose failure mode "
    "it can tolerate. We describe a live ingestion path that posts extracted "
    "photometry into SkyPortal in real time and release the pipeline as an "
    "installable package."
)
pa = doc.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pa.paragraph_format.left_indent = Inches(0.3)
pa.paragraph_format.right_indent = Inches(0.3)
ra = pa.add_run(abstract); ra.font.size = Pt(9)

kw = doc.add_paragraph(); kw.paragraph_format.left_indent = Inches(0.3)
kw.paragraph_format.right_indent = Inches(0.3)
kw.paragraph_format.space_before = Pt(4)
kk = kw.add_run("Unified Astronomy Thesaurus concepts: Time domain astronomy (2109); "
                "Transient sources (1851); Gamma-ray bursts (629); Astronomy data "
                "analysis (1858); Astronomy databases (83)")
kk.font.size, kk.italic = Pt(9), True

# ==================================================================
# switch to TWO columns for the body
# ==================================================================
body = doc.add_section(WD_SECTION.CONTINUOUS)
_cols(body, 2)


def fullwidth_start():
    s = doc.add_section(WD_SECTION.CONTINUOUS); _cols(s, 1); return s


def twocol_resume():
    s = doc.add_section(WD_SECTION.CONTINUOUS); _cols(s, 2); return s


# ---------------- 1. Introduction ----------------
heading(1, "Introduction")
para("Time-domain and multimessenger astronomy generate a continuous stream of "
     "human-written observation reports. NASA’s General Coordinates Network (GCN) "
     "distributes two products: machine-generated Notices, and free-text Circulars "
     "in which observers report follow-up of high-energy and multimessenger "
     "transients. The Circulars archive spans nearly three decades and exceeds "
     "40,500 documents, of which roughly 18,600 report optical observations. Their "
     "prose format — magnitudes, filters, redshifts, and telescopes embedded in "
     "natural language — makes bulk ingestion into structured systems such as "
     "SkyPortal (van der Walt et al. 2019) laborious.", before=1)
p = para("S25 established the modern treatment of this archive with large language "
         "models, contributing a neural topic-modeling pipeline, a contrastively "
         "fine-tuned classifier of observation wave bands, and a Mistral-based system "
         "for extracting gamma-ray burst redshifts, reporting 97.2% accuracy on "
         "redshift-containing Circulars retrieved by retrieval-augmented generation. "
         "We build directly on that work.")
footnote(p, "S25 code and data: https://github.com/GCN-Circulars-Analysis")
para("The present paper isolates a single design variable that S25 did not vary: the "
     "channel through which the language model’s output is obtained. Where S25 "
     "generate free-form text and parse it after the fact, we serve the identical "
     "model with grammar-constrained decoding, in which the sampler is masked at "
     "every step to tokens that continue a schema-valid string. Because the model is "
     "held fixed, the comparison measures the extraction harness rather than model "
     "capability. Two further contributions accompany this result: a transparent "
     "regular-expression baseline that establishes a floor on every field, and a "
     "live path that posts extracted photometry into SkyPortal as Circulars arrive.")
p = para("Circex is released as an installable Python package with a reproducible "
         "evaluation harness.")
footnote(p, "pip install circex — https://pypi.org/project/circex/")

# ---------------- 2. Schema + regex ----------------
heading(2, "Structured Representation and Regular-Expression Baseline")
subheading("2.1", "Methods")
para("2.1.1. Output schema. — Every extractor emits one Pydantic v2 model, "
     "CircularExtraction, whose JSON Schema derives from and extends "
     "nasa-gcn/gcn-schema. The model carries event identity, sky localization, a list "
     "of photometry points (filter, magnitude, uncertainty, magnitude system, "
     "limiting magnitude, and observation epoch), classification, redshift, and an "
     "optional provenance map that grounds each extracted value at a (start, end) "
     "character range in the source text for audit. A single output contract is what "
     "makes the extractors of Sections 2 and 3 commensurable.", before=1)
para("2.1.2. Regular-expression baseline. — Six sub-parsers cover event names, "
     "coordinates, redshift, photometry (single-mention, pipe-delimited, and "
     "fixed-width tables), classification, and time offsets. Each records the "
     "character span it consumed, populating the provenance map. The baseline is "
     "deterministic and transparent: it fabricates nothing, and its errors are "
     "inspectable.")
para("2.1.3. Evaluation. — We score against the first 500 rows of the 13,593-row "
     "Swift-validated redshift table released by S25, using a null-aware comparator "
     "with per-field precision, recall, and F1. Numeric fields use tolerances; list "
     "fields use greedy matching. Both our extractors and the S25 predictions are "
     "scored against the same Actual columns, so all denominators are identical.")
subheading("2.2", "Results")
para("On the three fields with non-zero gold support (Table 1), the regex baseline "
     "reaches F1 = 0.869 on event names and 0.862 on redshift, exceeding the "
     "published Mistral predictions by +0.020 and +0.171 respectively. That a "
     "transparent baseline already matches a language model on these fields is not a "
     "criticism of S25 but a calibration: it establishes the floor above which the "
     "constrained-decoding gain of Section 3 is measured.", before=1)

# ---- Table 1 (full width) ----
fullwidth_start()
para("Table 1", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=9.5, first_indent=0, before=6)
para("Per-field F1 on 500 rows of the S25 Swift-validated gold set. Support is the "
     "number of non-null gold values (TP + FN). The regex baseline does not attempt "
     "telescope-name extraction.", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True,
     size=8.5, first_indent=0, after=4)
tbl = doc.add_table(rows=1, cols=5); tbl.style = "Table Grid"
tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr = ["Field", "Support", "Regex (this work)", "Mistral-7B, constrained (this work)",
       "Mistral-7B (S25)"]
for c, h in zip(tbl.rows[0].cells, hdr):
    rp = c.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rp.add_run(h); rr.bold, rr.font.size = True, Pt(8.5)
rows = [
    ("Event name (GRB number)", "400", "0.869", "0.767", "0.849"),
    ("Redshift value", "383", "0.862", "0.935", "0.690"),
    ("Telescope name", "400", "n/a", "0.094", "0.098"),
]
for row in rows:
    cells = tbl.add_row().cells
    for i, val in enumerate(row):
        cp = cells[i].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        rr = cp.add_run(val); rr.font.size = Pt(8.5)
        if val in ("0.935", "0.869"):
            rr.bold = True
twocol_resume()

# ---------------- 3. Constrained LLM ----------------
heading(3, "Schema-Constrained Language Model Extraction")
subheading("3.1", "Methods")
para("3.1.1. Model and serving. — We use the same Mistral-7B-Instruct-v0.2 model "
     "as S25 (Jiang et al. 2023), served locally through a llama.cpp server exposing "
     "an OpenAI-compatible interface on a single GPU. The prompt — an extraction "
     "policy and few-shot examples — is shared across all backends; only the output "
     "channel varies.", before=1)
para("3.1.2. Two output configurations. — In the first, resembling the S25 harness, "
     "the schema is described in the prompt and the model is asked to comply in JSON "
     "mode; output is parsed afterward, with a single repair retry and a discard path "
     "on failure. This is a post-hoc constraint: nothing prevents malformed output, "
     "and when it occurs the extraction is lost. In the second, the schema is passed "
     "as a grammar (response_format: json_schema), so at each decoding step the "
     "sampler is restricted to tokens that can continue a schema-valid string "
     "(Willard & Louf 2023). The model cannot emit an unparseable response; no repair "
     "or discard path is required. This is the configuration reported as "
     "“constrained” throughout.")
para("3.1.3. Grammar construction. — The grammar is not the full schema but a "
     "bounded derivation of it. Supplying the complete CircularExtraction schema "
     "produced generations exceeding a five-minute timeout: an unbounded array lets "
     "the model emit photometry rows indefinitely; an 18-field object is emitted in "
     "full, mostly null, per row; and an unbounded string invites rambling within one "
     "value. Our grammar exposes only the scored fields, caps array lengths, prunes "
     "each object to the fields the model should produce (PhotometryExt 18→7; "
     "Localization 9→2), and caps string lengths — 72% smaller than the schema "
     "from which it derives (Appendix A). Fields Circex derives itself (canonical "
     "bandpass, epoch, detection flag, taxonomy path) are withheld from the model.")
subheading("3.2", "Results")
para("3.2.1. Redshift extraction. — The constrained Mistral-7B attains F1 = 0.935, "
     "exceeding both the regex baseline (0.862) and the identical model as published "
     "(0.690; Δ = +0.245). The mechanism is legible in the precision/recall "
     "decomposition (Table 2). Precision is comparable across the two Mistral "
     "configurations (0.903 and 0.932); the entire gap is in recall, which rises from "
     "0.559 to 0.937. The published pipeline is not extracting redshifts incorrectly "
     "— it fails to emit them, because free-form generation from a small model yields "
     "output that cannot be parsed, is discarded, and is scored as a false negative "
     "on every field the Circular in fact contained. Constrained decoding removes "
     "that failure mode by construction, and the model’s latent ability, previously "
     "masked by an unreliable serialization step, is measured directly.", before=1)

# ---- Table 2 ----
para("Table 2", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_indent=0, before=6)
para("Redshift extraction decomposed (383 gold values).",
     align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=8.5, first_indent=0, after=4)
t2 = doc.add_table(rows=1, cols=4); t2.style = "Table Grid"
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for c, h in zip(t2.rows[0].cells, ["Extractor", "Precision", "Recall", "F1"]):
    rp = c.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rp.add_run(h); rr.bold, rr.font.size = True, Pt(8.5)
for row in [("Regex", "0.870", "0.854", "0.862"),
            ("Mistral-7B (constrained)", "0.932", "0.937", "0.935"),
            ("Mistral-7B (S25)", "0.903", "0.559", "0.690")]:
    cells = t2.add_row().cells
    for i, val in enumerate(row):
        cp = cells[i].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        rr = cp.add_run(val); rr.font.size = Pt(8.5)

para("We state the finding plainly: the reported weakness of open 7B models on this "
     "task is substantially an artifact of the extraction harness rather than of the "
     "model. The same weights that score 0.690 score 0.935 when the decoder is "
     "constrained — a gain of +0.245 F1 with no fine-tuning, no retrieval, and no "
     "larger model. This also reconciles an apparent tension with the 97.2% headline "
     "of S25: that figure is conditional on Circulars known to contain a redshift, "
     "after retrieval, whereas the F1 here is unconditioned over the released "
     "predictions; both are scored against the same denominator as those predictions.",
     before=4)
para("3.2.2. Event names and the hybrid. — The converse is equally clean. On event "
     "names the ordering inverts: regex (0.869) leads, the published Mistral (0.849) "
     "follows, and the constrained Mistral is worst (0.767). Its error decomposition "
     "(344 true positives, 153 false positives, 56 false negatives, versus regex’s "
     "388/105/12) shows it both misses and invents more designations. A GCN event "
     "name is a lexically regular object — GRB followed by a six-digit date and an "
     "optional letter — which is precisely a regular expression’s task; constrained "
     "decoding guarantees well-formed output, not pattern-matching skill. Routing "
     "each field to the extractor that wins it — regex for event names, constrained "
     "Mistral for redshift — yields a hybrid that dominates every single extractor on "
     "every field with gold support, and is how the SkyPortal path of Section 4 is "
     "configured.")
para("3.2.3. Cost and latency. — The constrained extractor runs at a median 1.5 s "
     "per Circular (p95 4.6 s) on one GPU, against 0.6 ms for regex. The 500-row "
     "evaluation completes in roughly fifteen minutes at zero marginal token cost, "
     "the model being served locally. At that throughput the full optical archive is "
     "a single-day backfill and live per-Circular ingestion is not throughput-bound.")

# ---- Figure 1 (full width) ----
fullwidth_start()
if FIG.exists():
    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run().add_picture(str(FIG), width=Inches(6.6))
para("Figure 1.", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_indent=0, before=2)
para("Per-field F1 for the regex baseline, the grammar-constrained Mistral-7B "
     "extractor, and the published Mistral-7B predictions of S25 (top), and ΔF1 "
     "against the published baseline (bottom). Hatched bars denote a non-extracting "
     "extractor or zero gold support.", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True,
     size=8.5, first_indent=0, after=4)
twocol_resume()

# ---------------- 4. Deployment ----------------
heading(4, "Deployment: Live Ingestion into SkyPortal")
subheading("4.1", "Methods")
para("4.1.1. Event aggregation. — A transient is described across many Circulars: a "
     "discovery Circular carries the position, follow-ups carry the light curve. "
     "Circex reconstructs an event by walking the GCN cross-reference graph the "
     "extractor already recovers, then fuses the Circulars into one source — "
     "position preferring the refined optical counterpart over a coarse trigger box, "
     "photometry unioned and deduplicated with an epoch-tolerant key so the same "
     "observation reported at different epochs is not double-counted.", before=1)
p = para("4.1.2. The consumer service. — A consumer subscribes to the gcn.circulars "
         "Kafka stream; for each incoming Circular it aggregates the event and posts "
         "the source, photometry, and classification into SkyPortal idempotently, "
         "tolerating individual failures without interrupting the stream.")
footnote(p, "SkyPortal: https://skyportal.io")
subheading("4.2", "Validation")
para("The path was validated end-to-end against the live GCN stream and a production "
     "SkyPortal instance: a real gamma-ray burst was reconstructed from its Circulars "
     "and its multi-telescope, multi-band light curve posted, with idempotency "
     "confirmed on re-processing. The classification field is supplied by a lightweight "
     "Naive Bayes classifier trained on labels harvested from the archive, which "
     "abstains on Circulars carrying no supernova type — the failure mode that makes "
     "a naive regex classifier over-fire.", before=1)

# ---------------- 5. Discussion ----------------
heading(5, "Discussion and Conclusion")
para("Holding the model fixed and varying only the output channel isolates a result "
     "that we believe is favorable to the thesis of S25 rather than a correction to "
     "it: open 7B models are more capable at structured scientific extraction than "
     "free-form-generation numbers indicate, once the serialization bottleneck is "
     "removed. The +0.245 F1 recovered on redshift is obtained with the released "
     "weights, no fine-tuning, and no retrieval. Because the model is served locally, "
     "the marginal cost per Circular is zero and no text leaves the institution — "
     "removing both the recurring cost and the data-governance objection that would "
     "attach to a commercial API in a continuous observatory pipeline.", before=1)
subheading("5.1", "Challenges and Limitations")
para("First, the redshift result is a lower bound on the constrained model’s "
     "advantage: the fields on which a language model should separate most sharply "
     "from regex — multi-row photometry tables, in-prose classification, unlabeled "
     "coordinates — are precisely those S25 did not extract, so no gold exists to "
     "score them. A hand-labeled set covering these fields is the natural next "
     "experiment. Second, constrained decoding must be engineered, not merely "
     "enabled: a naive schema-to-grammar conversion of a realistic Pydantic model "
     "does not work, and every unbounded degree of freedom becomes a way for a small "
     "model to exhaust its token budget (Appendix A). Third, telescope-name "
     "extraction remains weak for all extractors, reflecting a normalization gap "
     "between formal catalog codes and prose mentions rather than a failure of "
     "extraction as such.")

# ---------------- Acknowledgments ----------------
heading(None, "Acknowledgments")
para("We thank Sushant Sharma Chaudhary for deploying the Mistral-7B llama.cpp "
     "server used for the constrained-decoding experiments and for pointing us to "
     "grammar-constrained decoding through the server’s response_format interface, "
     "the mechanism on which the central result rests. The authors acknowledge the "
     "Minnesota Supercomputing Institute (MSI) at the University of Minnesota for "
     "providing the GPU resources that contributed to the results reported here. This "
     "work builds directly on S25 and uses their released Circulars archive, topic "
     "labels, and Swift-validated redshift table, together with the "
     "skyportal/timedomain-taxonomy controlled vocabulary. Portions of the SQLite/FTS "
     "indexer and the GCN poller were ported from the sjhend03/GCNMCP prototype with "
     "attribution.", before=1, size=9)

# ---------------- References ----------------
heading(None, "References")
refs = [
    "Abbott, B. P., Abbott, R., Abbott, T. D., et al. 2017, ApJL, 848, L12",
    "Gao, Y., Xiong, Y., Gao, X., et al. 2023, arXiv:2312.10997",
    "Gehrels, N., Chincarini, G., Giommi, P., et al. 2004, ApJ, 611, 1005",
    "Jiang, A. Q., Sablayrolles, A., Mensch, A., et al. 2023, arXiv:2310.06825",
    "Lewis, P., Perez, E., Piktus, A., et al. 2020, in Advances in Neural Information "
    "Processing Systems, Vol. 33 (Red Hook, NY: Curran Associates), 9459",
    "Sharma, V., Agarwala, R., Racusin, J. L., et al. 2026, ApJS, 283, 30",
    "van der Walt, S., Crellin-Quick, A., & Bloom, J. S. 2019, JOSS, 4, 1247",
    "Willard, B. T., & Louf, R. 2023, arXiv:2307.09702",
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rr = p.add_run(r); rr.font.size = Pt(8.5)

# ---------------- Appendix A ----------------
heading(None, "Appendix A")
para("Grammar Bounding", bold=True, first_indent=0, before=2)
para("The grammar passed to the llama.cpp server is a bounded derivation of the "
     "CircularExtraction JSON Schema. Table 3 records the reduction. Beyond field "
     "pruning, array lengths are capped (maxItems: photometry 15, time_offsets 10), "
     "string lengths are capped (maxLength 128), and a hard generation limit "
     "(max_tokens 2048) provides a final backstop. Together these bound the grammar "
     "in all three dimensions — arrays, objects, and strings — each of which a "
     "small model will otherwise exploit until it exhausts the token budget.",
     before=1, size=9)
para("Table 3", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_indent=0, before=6)
para("Schema reduction for grammar-constrained decoding.",
     align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=8.5, first_indent=0, after=4)
t3 = doc.add_table(rows=1, cols=3); t3.style = "Table Grid"
for c, h in zip(t3.rows[0].cells, ["Quantity", "Full schema", "Grammar (lean)"]):
    rp = c.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rp.add_run(h); rr.bold, rr.font.size = True, Pt(8.5)
for row in [("Serialized size (chars)", "21,350", "5,786"),
            ("Schema definitions ($defs)", "13", "6"),
            ("PhotometryExt fields", "18", "7"),
            ("Localization fields", "9", "2")]:
    cells = t3.add_row().cells
    for i, val in enumerate(row):
        cp = cells[i].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run(val).font.size = Pt(8.5)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print("wrote", OUT)
